# -*- coding: utf-8 -*-
"""
Module de gestion de documents pour Cypher
Système RAG (Retrieval-Augmented Generation) avec ChromaDB
🚀 OPTIMISÉ : Lazy Loading des dépendances lourdes
"""

import os
from core.paths import get_rag_db_dir
from core.logger import get_logger

logger = get_logger("document_manager")


def document_manager(
    action: str, 
    query: str | None = None, 
    source_folder: str | None = None, 
    source_file: str | None = None
) -> str:
    """
    Système RAG (Chat with Data) : Indexe et interroge vos documents locaux (PDF, MD, DOCX, TXT).
    Action: 'index' (scanner un dossier) ou 'search' (poser une question).
    
    🚀 OPTIMISATION : Les dépendances lourdes (chromadb, fitz, docx) sont importées
    uniquement au premier appel, accélérant le démarrage de Cypher de 3-5 secondes.
    """
    # 🚀 LAZY IMPORT : Imports lourds seulement quand nécessaire
    import fitz  
    import docx
    import chromadb
    from chromadb.utils import embedding_functions

    # Chemin de la base de données vectorielle (mémoire documentaire)
    DB_PATH = str(get_rag_db_dir())
    
    # Initialisation de ChromaDB (Persistant sur le disque)
    client = chromadb.PersistentClient(path=DB_PATH)
    
    # Fonction d'embedding par défaut (utilise un modèle léger local)
    # Note: Au premier lancement, ça téléchargera un petit modèle (~80MB)
    emb_fn = embedding_functions.SentenceTransformerEmbeddingFunction(model_name="all-MiniLM-L6-v2")
    
    # Création/Récupération de la collection
    collection = client.get_or_create_collection(name="cypher_knowledge", embedding_function=emb_fn)

    action = action.lower()

    # --- 1. INDEXATION (APPRENDRE) ---
    if action == "index":
        if not source_folder:
            return "Quel dossier dois-je lire et apprendre ?"
        
        # Correction chemin OneDrive si tilde
        if "~" in source_folder:
            source_folder = os.path.expanduser(source_folder)
        
        if not os.path.exists(source_folder):
            return f"Le dossier '{source_folder}' n'existe pas."

        logger.info(f"Indexation en cours de : {source_folder} ...")
        
        files_processed = 0
        chunks_added = 0
        
        # Formats supportés
        supported_ext = ['.pdf', '.md', '.txt', '.docx', '.py']

        for root, dirs, files in os.walk(source_folder):
            for file in files:
                ext = os.path.splitext(file)[1].lower()
                if ext in supported_ext:
                    filepath = os.path.join(root, file)
                    text_content = ""
                    
                    try:
                        # Extraction du texte selon le format
                        if ext == '.pdf':
                            try:
                                with fitz.open(filepath) as doc:
                                    for page in doc:
                                        text_content += page.get_text() + "\n"
                            except fitz.fitz.FileDataError:
                                logger.warning(f"PDF corrompu : {file}")
                                continue
                            except Exception as e:
                                logger.warning(f"Erreur PDF {file}: {e}")
                                continue
                        elif ext == '.docx':
                            doc = docx.Document(filepath)
                            text_content = "\n".join([p.text for p in doc.paragraphs])
                        else:  # MD, TXT, PY
                            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                                text_content = f.read()
                        
                        if not text_content.strip():
                            continue

                        # Découpage en chunks (morceaux) de 1000 caractères
                        # Pour éviter de saturer le contexte
                        chunk_size = 1000
                        overlap = 100  # On garde 100 char du chunk précédent
                        chunks = []
                        for i in range(0, len(text_content), chunk_size - overlap):
                            chunks.append(text_content[i:i+chunk_size])
                        
                        # Préparation pour ChromaDB
                        ids = [f"{file}_{i}" for i in range(len(chunks))]
                        metadatas = [{"source": file, "path": filepath} for _ in chunks]
                        
                        # Ajout à la base (Upsert = met à jour si existe déjà)
                        collection.upsert(
                            documents=chunks,
                            ids=ids,
                            metadatas=metadatas
                        )
                        
                        files_processed += 1
                        chunks_added += len(chunks)
                        logger.info(f"Lu : {file} ({len(chunks)} fragments)")
                        
                    except Exception as e:
                        logger.warning(f"Erreur sur {file}: {e}")
                        continue

        return f"Indexation terminée ! J'ai lu {files_processed} fichiers et mémorisé {chunks_added} fragments de connaissances."

    # --- 2. RECHERCHE (RÉPONDRE) ---
    if action == "search":
        if not query:
            return "Quelle est votre question sur les documents ?"
        
        # On cherche les 5 morceaux les plus pertinents sémantiquement
        results = collection.query(
            query_texts=[query],
            n_results=5
        )
        
        if not results['documents'][0]:
            return "Je n'ai trouvé aucune information pertinente dans vos documents indexés."
        
        # Construction du contexte pour Gemini
        context_text = "Voici les informations extraites de vos documents :\n\n"
        for i, doc in enumerate(results['documents'][0]):
            source = results['metadatas'][0][i]['source']
            context_text += f"--- Source : {source} ---\n{doc}\n\n"
        
        return context_text

    # --- 3. RESET (OUBLIER) ---
    if action == "reset":
        client.delete_collection("cypher_knowledge")
        return "J'ai effacé toute ma base de connaissances documentaires."

    # --- 4. SUMMARY (RÉSUMER UN FICHIER) ---
    if action == "summary":
        if not source_file:
            return "Quel fichier dois-je résumer, Monsieur ?"

        # Expansion du chemin si besoin (~, etc.)
        if "~" in source_file:
            source_file = os.path.expanduser(source_file)

        if not os.path.exists(source_file):
            return f"Le fichier '{source_file}' n'existe pas, Monsieur."

        ext = os.path.splitext(source_file)[1].lower()
        text_content = ""

        try:
            if ext == ".pdf":
                try:
                    with fitz.open(source_file) as doc:
                        for page in doc:
                            text_content += page.get_text() + "\n"
                except Exception as e:
                    return f"Impossible de lire le PDF, Monsieur : {e}"

            elif ext == ".docx":
                try:
                    d = docx.Document(source_file)
                    text_content = "\n".join([p.text for p in d.paragraphs])
                except Exception as e:
                    return f"Impossible de lire le document Word, Monsieur : {e}"

            elif ext in [".md", ".txt", ".py"]:
                with open(source_file, "r", encoding="utf-8", errors="ignore") as f:
                    text_content = f.read()

            else:
                return f"Je ne sais pas encore résumer les fichiers de type '{ext}', Monsieur."

            if not text_content.strip():
                return "Le fichier est vide ou illisible, Monsieur."

            # On limite la taille renvoyée pour ne pas exploser le contexte de Gemini
            max_chars = 4000
            snippet = text_content[:max_chars]
            if len(text_content) > max_chars:
                snippet += "\n\n[Texte tronqué pour le résumé, Monsieur.]"

            # L'idée : tu renvoies le texte brut, et Gemini se charge de faire un beau résumé
            return (
                f"Voici le contenu brut extrait de '{os.path.basename(source_file)}', "
                f"prêt à être résumé :\n\n{snippet}"
            )

        except Exception as e:
            return f"Erreur lors de la lecture du fichier, Monsieur : {e}"

    return "Action RAG inconnue."

